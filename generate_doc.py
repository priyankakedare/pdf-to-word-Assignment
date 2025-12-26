from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL


def generate_word(output_path="Mediation_Application_Form.docx"):
    doc = Document()

    # ---------------- Page Margins ----------------
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    # ---------------- Centered Headings ----------------
    def center(text, bold=False, size=12):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(size)

    center("FORM ‘A’", True, 14)
    center("MEDIATION APPLICATION FORM", True, 14)
    center("[REFER RULE 3(1)]", False, 11)
    center("Mumbai District Legal Services Authority")
    center("City Civil Court, Mumbai")
    doc.add_paragraph("")

    # ---------------- Table (4 columns like PDF) ----------------
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"

    widths = [Inches(0.5), Inches(1.3), Inches(2.4), Inches(2.8)]
    for row in table.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = w

    # ---------------- Header Row ----------------
    hdr = table.rows[0].cells
    hdr[0].merge(hdr[3])
    hdr[0].paragraphs[0].add_run("DETAILS OF PARTIES:").bold = True

    def add_row(c1="", c2="", c3="", c4=""):
        r = table.add_row().cells
        r[0].text = c1
        r[1].text = c2
        r[2].text = c3
        r[3].text = c4

    # ---------------- Applicant (Index 1) ----------------
    start_app = len(table.rows)

    add_row("", "Name of", "Applicant", "{{client_name}}")
    add_row("", "Address", "REGISTERED ADDRESS:", "{{branch_address}}")
    add_row("", "", "CORRESPONDENCE BRANCH ADDRESS:", "{{branch_address}}")
    add_row("", "", "Telephone No.", "{{mobile}}")
    add_row("", "", "Mobile No.", "")
    add_row("", "", "Email ID", "info@kslegal.co.in")

    # Merge index column (1)
    idx = table.rows[start_app].cells[0]
    for i in range(start_app + 1, start_app + 6):
        idx.merge(table.rows[i].cells[0])
    idx.text = "1"
    idx.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Merge Address label
    addr = table.rows[start_app + 1].cells[1]
    addr.merge(table.rows[start_app + 2].cells[1])
    addr.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # ---------------- Opposite Party (Index 2) ----------------
    r = table.add_row().cells
    r[0].merge(r[3])
    r[0].paragraphs[0].add_run(
        "2   Name, Address and Contact details of Opposite Party:"
    ).bold = True

    start_opp = len(table.rows)

    add_row("", "Name", "", "{{customer_name}}")
    add_row("", "Address", "REGISTERED ADDRESS:", "________________")
    add_row("", "", "CORRESPONDENCE ADDRESS:", "________________")
    add_row("", "", "Telephone No.", "")
    add_row("", "", "Mobile No.", "")
    add_row("", "", "Email ID", "")

    # Merge index column (2)
    idx = table.rows[start_opp].cells[0]
    for i in range(start_opp + 1, start_opp + 6):
        idx.merge(table.rows[i].cells[0])
    idx.text = "2"
    idx.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Merge Address label
    addr = table.rows[start_opp + 1].cells[1]
    addr.merge(table.rows[start_opp + 2].cells[1])
    addr.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # ---------------- Dispute Section ----------------
    r = table.add_row().cells
    r[0].merge(r[3])
    r[0].paragraphs[0].add_run("DETAILS OF DISPUTE:").bold = True

    r = table.add_row().cells
    r[0].merge(r[3])
    r[0].paragraphs[0].add_run(
        "THE COMM. COURTS (PRE-INSTITUTION SETTLEMENT) RULES, 2018"
    ).bold = True

    r = table.add_row().cells
    r[0].merge(r[3])
    r[0].paragraphs[0].add_run(
        "Nature of disputes as per section 2(1)(c) of the Commercial Courts Act, 2015 (4 of 2016):"
    ).bold = True

    # ---------------- Save ----------------
    doc.save(output_path)
    return output_path


if __name__ == "__main__":
    generate_word()
